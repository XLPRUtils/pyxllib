#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# @Author : 陈坤泽
# @Email  : 877362867@qq.com
# @Date   : 2021/08/31 09:56

from pyxllib.prog.pupil import check_install_package

check_install_package('win32com', 'pypiwin32')
check_install_package('docx', 'python-docx')

import json
import os
import re

import pythoncom
from win32com.client import constants
import win32com.client as win32
import docx
import docx.table
import docx.enum

from pyxllib.prog.pupil import DictTool, inject_members, run_once
from pyxllib.text.pupil import strwidth
from pyxllib.debug.specialist import File, Dir, get_etag, browser


def __docx():
    """ python-docx 相关封装
    """
    pass


class DocxTools:
    @classmethod
    def to_pdf(cls, docx_file, pdf_file=None):
        check_install_package('docx2pdf')  # 安装不成功的时候可以考虑加参数：--user
        import docx2pdf

        if pdf_file is None:
            pdf_file = docx_file.with_suffix('.pdf')

        docx2pdf.convert(str(docx_file), str(pdf_file))
        return pdf_file

    @classmethod
    def merge(cls, master_file, toc, *, outline='demote'):
        """ 合并多份docx文件

        :param master_file: 要合并到哪个主文件
            注意如果这个文件已存在，会被替换，重置
        :param toc: 类似fitz的table of contents，用一个n*3的表格表示新文档的格式
            list，每个元素三列：标题级别，标题名称，(可选)对应文件内容
        :param outline: 原来每份子文档里的标题，插入到新文档中的处理规则
            demote：降级
            clear：清除

        这个功能还有些局限性，后面要扩展鲁棒性
        TODO 增加一个支持将原文档标题降级的功能，降到toc之后
        """
        app = XlWin32WordApplication.get_app()

        master_doc = app.new_doc(master_file)
        for item in toc:
            lvl, title, file = item

            # 1 加一个标题
            r = master_doc.Paragraphs.Add().Range
            r.InsertBefore(title)
            r.Style = master_doc.Styles(f'标题 {lvl}')

            # 2 拷贝完整的内容
            if file:
                file = File(file)
                member_doc = app.open_doc(file)
                member_doc.Activate()

                # 处理原来文档的目录级别
                if outline == 'demote':
                    member_doc.outline_demote(lvl)
                elif outline == 'clear':
                    # 降10级就相当于清除所有原来的标题
                    member_doc.outline_demote(10)

                app.Selection.WholeStory()
                app.Selection.Copy()
                master_doc.Activate()
                app.Selection.EndKey(Unit=app.wd('Story'))  # 跳到文档末尾
                app.Selection.Paste()
                member_doc.Close()
        master_doc.save()
        master_doc.Close(True)


class Document:
    """ 这个库写英文文档还不错。但不能做中文，字体会错乱。
    """

    def __init__(self, docx_file=None):
        """
        Args:
            docx_file:
                已有的word文件路径：打开
                还没创建的word文件路径：在个别功能需要的时候，会自动创建
                None：在临时文件夹生成一个默认的word文件
        """
        if docx_file is None:
            self.docx_file = File(..., Dir.TEMP, suffix='.docx')
        else:
            self.docx_file = File(docx_file)
        if self.docx_file:
            self.doc = docx.Document(str(docx_file))
        else:
            self.doc = docx.Document()

    def write(self):
        Dir(self.docx_file.parent).ensure_dir()
        self.doc.save(str(self.docx_file))

    def to_pdf(self, pdf_file=None):
        self.write()
        pdf_file = DocxTools.to_pdf(self.docx_file, pdf_file)
        return pdf_file

    def to_fitzdoc(self):
        """ 获得 fitz的pdf文档对象
        :return: FitzDoc对象
        """
        from pyxllib.file.pdflib import FitzDoc
        pdf_file = self.to_pdf()
        doc = FitzDoc(pdf_file)
        return doc

    def to_images(self, file_fmt='{filestem}_{number}.png', *args, scale=1, **kwargs):
        doc = self.to_fitzdoc()
        files = doc.to_images(doc.src_file.parent, file_fmt, *args, scale=scale, **kwargs)
        return files

    def browser(self):
        """ 转pdf，使用浏览器的查看效果
        """
        pdf_file = self.to_pdf()
        browser(pdf_file)

    def display(self):
        """ 转图片，使用jupyter环境的查看效果
        """
        from pyxllib.cv.expert import PilImg

        # 转图片，并且裁剪，加边框输出
        doc = self.to_fitzdoc()
        for i in range(doc.page_count):
            page = doc.load_page(i)
            print('= ' * 10 + f' {page} ' + '= ' * 10)
            img: PilImg = page.get_pil_image()
            img.trim(border=5).plot_border().display()
            del page

    def to_labelmes(self, dst_dir=None, file_fmt='{filestem}_{number}.png', *, views=(0, 0, 1, 0), scale=1,
                    advance=False, indent=None):
        """ 转labelme格式查看

        本质是把docx转成pdf，利用pdf的解析生成labelme格式的标准框查看

        :param views: 详见to_labelmes的描述
            各位依次代表是否显示对应细粒度的标注：blocks、lines、spans、chars
        :param bool|dict advance: 是否开启“高级”功能，开启后能获得下划线等属性，但速度会慢很多
            源生的fitz pdf解析是处理不了下划线的，开启高级功能后，有办法通过特殊手段实现下划线的解析
            默认会修正目前已知的下划线、颜色偏差问题
            dict类型：以后功能多了，可能要支持自定义搭配，比如只复原unberline，但不管颜色偏差
        """
        from pyxlpr.data.labelme import LabelmeDict

        # 1 转成图片，及json标注
        doc = self.to_fitzdoc()
        imfiles = doc.to_images(dst_dir, file_fmt, scale=scale)

        # 2 高级功能
        def is_color(x):
            return x and sum(x)

        def to_labelmes_advance():
            m = 50  # 匹配run时，上文关联的文字长度，越长越严格

            # 1 将带有下划线的run对象，使用特殊的hash规则存储起来
            content = []  # 使用已遍历到的文本内容作为hash值
            elements = {}
            for p in self.paragraphs:
                for r in p.runs:
                    # 要去掉空格，不然可能对不上。试过strip不行。分段会不太一样，还是把所有空格删了靠谱。
                    content.append(re.sub(r'\s+', '', r.text))
                    if r.underline or is_color(r.font.color.rgb):  # 对有下划线、颜色的对象进行存储
                        # print(r.text + ',', r.underline, r.font.color.rgb, ''.join(content))
                        etag = get_etag(''.join(content)[-m:])  # 全部字符都判断的话太严格了，所以减小区间~
                        elements[etag] = r

            # 2 检查json标注中为span的，其哈希规则是否有对应，则要单独设置扩展属性
            content = ''
            for i, file in enumerate(imfiles):
                page = doc.load_page(i)
                lmdict = LabelmeDict.gen_data(file)
                lmdict['shapes'] = page.get_labelme_shapes('dict', views=views, scale=scale)

                for sp in lmdict['shapes']:
                    attrs = DictTool.json_loads(sp['label'], 'label')
                    if attrs['category_name'] == 'span':
                        content += re.sub(r'\s+', '', attrs['text'])
                        etag = get_etag(content[-m:])
                        # print(content)
                        if etag in elements:
                            # print(content)
                            r = elements[etag]  # 对应的原run对象
                            attrs = DictTool.json_loads(sp['label'])
                            x = r.underline
                            if x:
                                attrs['underline'] = int(x)
                            x = r.font.color.rgb
                            if is_color(x):
                                attrs['color'] = list(x)
                            sp['label'] = json.dumps(attrs)
                file.with_suffix('.json').write(lmdict, indent=indent)

        # 3 获得json
        if advance:
            to_labelmes_advance()
        else:
            doc.to_labelmes(imfiles, views=views, scale=scale)

    def __getattr__(self, item):
        # 属性：
        # core_properties
        # element
        # inline_shapes
        # paragraphs
        # part
        # sections
        # settings
        # styles
        # tables
        # 方法：
        # add_heading
        # add_page_break
        # add_paragraph
        # add_picture
        # add_section
        # add_table
        # save
        return getattr(self.doc, item)


class XlDocxTable(docx.table.Table):
    def merge_samevalue_in_col(self, col, start_row=1):
        """ 定义合并单元格的函数

        :param col: 需要处理数据的列，0开始编号
        :param start_row: 起始行，即表格中开始比对数据的行（其实标题排不排除一般无所谓~默认是排除了）
        """

        def merge_cells(start, end):
            if end > start:
                c = self.cell(start, col)
                c.merge(self.cell(end, col))
                c.text = self.cell(start, col).text.strip()
                c.vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.CENTER
                c.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

        ref, start = None, start_row
        for i in range(start_row, len(self.rows)):
            v = self.cell(i, col).text
            if v != ref:
                merge_cells(start, i - 1)
                ref, start = v, i
            else:
                self.cell(i, col).text = ''
        merge_cells(start, i)


inject_members(XlDocxTable, docx.table.Table)


def __win32_word():
    """ 使用win32com调用word

    vba的文档：示例代码更多，vba语法也更熟悉，但显示的功能更不全
        https://docs.microsoft.com/en-us/office/vba/api/word.saveas2
    .net的文档：功能显示更全，应该是所有COM接口都有但示例代码更少、更不熟系
        https://docs.microsoft.com/en-us/dotnet/api/microsoft.office.interop.word.documentclass.saveas2?view=word-pia
    """
    pass


@run_once
def inject_win32word(app, recursion_inject=False):
    """ 给win32的word com接口添加功能

    :param app: win32的类是临时生成的，需要给一个参考对象，才方便type(word)算出类型
    :param recursion_inject: 是否递归，对目前有的各种子类扩展功能都绑定上
        默认关闭，如果影响到性能，可以关闭，后面运行中需要时手动设定inject注入
        开启，能方便业务层开发

        之前有想过可以生成doc里的时候再inject这些对象，但如果是批量处理脚本，每次建立doc都判断我觉得也麻烦
        长痛不如短痛，建立app的时候就把所有对象inject更方便
    """
    inject_members(XlWin32WordApplication, type(app), ignore_case=True)
    if recursion_inject:
        # 建一个临时文件，把各种需要绑定的对象都生成绑定一遍
        # 确保初始化稍微慢点，但后面就方便了
        doc = app.Documents.Add()
        inject_members(XlWin32WordDocument, type(doc), ignore_case=True)

        doc.Activate()
        rng = doc.Range()  # 全空的文档，有区间[0,1)
        inject_members(XlWin32WordRange, type(rng), ignore_case=True)

        doc.Hyperlinks.Add(rng, 'url')  # 因为全空，这里会自动生成对应的明文url
        inject_members(XlWin32WordHyperlink, type(doc.Hyperlinks(1)), ignore_case=True)

        # 处理完关闭文档，不用保存
        doc.Close(False)


class XlWin32WordApplication:
    @classmethod
    def get_app(cls, app=None, *, visible=None, display_alerts=0, recursion_inject=True):
        """
        Args:
            app: 可以自定义在外部使用Dispatch、DispatchEx等形式给入已初始化好的app
            visible: 是否可见
            display_alerts: 是否关闭警告
            recursion_inject: 是否递归执行inject
        """
        # 1 get app
        name = 'WORD.APPLICATION'
        if app is None:
            try:
                app = win32.GetActiveObject(name)
            except pythoncom.com_error:
                pass
        if app is None:
            try:
                # 名称用大写，会比较兼容旧的word2013等版本
                # 尽量静态调度，才能获得 from win32com.client import constants 的常量
                app = win32.gencache.EnsureDispatch(name)
            except TypeError:
                # 实在不行，就用动态调度
                app = win32.dynamic.Dispatch(name)
            # 注：好像docx的默认打开程序也会有影响，最好默认都是office，不要被改成wps

        # 2 inject
        inject_win32word(app, recursion_inject=recursion_inject)

        if visible is not None:
            app.Visible = visible
        if display_alerts is not None:
            app.DisplayAlerts = display_alerts  # 不警告

        return app

    def check_close(self, outfile):
        """ 检查是否有指定名称的文件被打开，将其关闭，避免new_doc等操作出现问题
        """
        outfile = File(outfile)
        for x in self.Documents:
            # 有可能文件来自onedrive，这里用safe_init更合理
            if File.safe_init(x.Name, x.Path) == outfile:
                x.Close()

    def open_doc(self, file_name):
        """ 打开已有的文件
        """
        doc = self.Documents.Open(str(file_name))
        return doc

    def new_doc(self, file=None):
        """ 创建一个新的文件
        Args:
            file: 文件路径
                空：新建一个doc，到时候保存会默认到临时文件夹
                不存在的文件名：新建对应的空文件
                已存在的文件名：重置、覆盖一个新的空文件

        使用该函数，会自动执行XlWin32WordDocument扩展。
        """
        if file is None:
            file = File(..., Dir.TEMP, suffix='.docx')
        else:
            file = File(file)

        doc = self.Documents.Add()  # 创建新的word文档
        doc.save(file)
        return doc

    @classmethod
    def wd(cls, name, part=None):
        """ 输入字符串名称，获得对应的常量值

        :param name: 必须省略前缀wd。这个函数叫wd，就是帮忙省略掉前缀wd的意思
        :param part: 特定组别的枚举值，可以输入特殊的name模式来取值
        """
        if part is None:
            return getattr(constants, 'wd' + name)
        else:
            raise ValueError


class XlWin32WordDocument:
    def save(self, file_name=None, fmt=None, retain=False, **kwargs):
        """ 我自己简化的保存接口

        :param file_name: 保存到指定路径，如果带有后缀
        :param fmt: 毕竟是底层的com接口，不能做的太智能吧。连通过文件名后缀自动选择格式的功能都没有，要手动指定。
            为了方便，对这些进行智能自动处理，得到一个合理的save接口。
        :param retain: SaveAs2的机制：如果目标格式仍是word支持的，则doc会切换到目标文件。否则doc保留原文件对象。
            这里retain若打开，则会自动做切换，保留原文件对象
        :return: 跟retain有关，可能会"重置", outfile
            默认返回 outfile
            开启retain时，返回 outfile, doc
        """

        # 1 辅助函数
        def save_format(fmt):
            """ 枚举值映射，word保存类型的枚举

            >> _('.html')
            8
            >> _('Pdf')
            17
            >> _('wdFormatFilteredHTML')
            10
            """
            # 复杂格式可能无法完美支持所有功能。比如复杂的pdf无法使用SaveAs2实现，要用ExportAsFixedFormat。
            # 有些情况可能无法使用gencache，导致没有constants，所以默认可以直接映射到整数，避免使用到constants
            common = {'doc': 0,
                      'html': 8,
                      'txt': 2,
                      'docx': 16,
                      'pdf': 17}
            name = common.get(fmt.lower().lstrip('.'), fmt)
            if isinstance(name, int):
                return name
            else:
                return getattr(constants, 'wd' + name)

        # 2 确认要存储的文件格式
        if isinstance(fmt, str):
            fmt = fmt.lower().lstrip('.')
        elif file_name is not None:
            fmt = File(file_name).suffix[1:].lower()
        elif self.Path:
            fmt = os.path.splitext(self.Name)[1][1:].lower()
        else:
            fmt = 'docx'

        # 3 保存一份原始的文件路径
        origin_file = File(self.Name, self.Path) if self.Path else None

        # 4 如果有指定保存文件路径
        if file_name is not None:
            outfile = File(file_name)
            if outfile.suffix[1:].lower() != fmt:
                # 已有文件名，但这里指定的fmt不同于原文件，则认为是要另存为一个同名的不同格式文件
                outfile = File(outfile.stem, outfile.parent, suffix=fmt)
            self.SaveAs2(str(outfile), save_format(fmt), **kwargs)
        # 5 如果没指定保存文件路径
        else:
            if self.Path:
                outfile = File(self.Name, self.Path, suffix='.' + fmt)
                self.SaveAs2(str(outfile), save_format(outfile.suffix), **kwargs)
            else:
                etag = get_etag(self.Content)
                outfile = File(etag, Dir.TEMP, suffix=fmt)
                self.SaveAs2(str(outfile), save_format(fmt), **kwargs)

        # 6 是否恢复原doc
        cur_file = File(self.Name, self.Path)  # 当前文件不一定是目标文件f，如果是pdf等格式也不会切换过去
        if retain and origin_file and origin_file != cur_file:
            app = self.Application
            self.Close()
            self = app.open_doc(origin_file)

        # 7 返回值
        if retain:
            return outfile, self
        else:
            return outfile

    # 先暂时不开启 doc.chars
    # @staticmethod
    # def chars(doc):
    #     return doc.Range().chars

    @property
    def n_page(self):
        return self.ActiveWindow.Panes(1).Pages.Count

    def browser(self, file_name=None, fmt='html', retain=False):
        """ 这个函数可能会导致原doc指向对象被销毁，建议要不追返回值doc继续使用
        """
        res = self.save(file_name, fmt, retain=retain)

        if retain:
            outfile, self = res
        else:
            outfile = res

        browser(outfile)
        return self

    def add_section_size(self, factor=1):
        """ 显示每节长度的标记
        一般在这里计算比在html计算方便
        """
        from humanfriendly import format_size

        n = self.Paragraphs.Count
        style_names, text_lens = [], []
        for p in self.Paragraphs:
            style_names.append(str(p.Style))
            text_lens.append(strwidth(p.Range.Text))

        for i, p in enumerate(self.Paragraphs):
            name = style_names[i]
            if name.startswith('标题'):
                cumulate_size = 0
                for j in range(i + 1, n):
                    if style_names[j] != name:
                        cumulate_size += text_lens[j - 1]
                    else:
                        break
                if cumulate_size:
                    size = format_size(cumulate_size * factor).replace(' ', '').replace('bytes', 'B')
                    r = p.Range
                    self.Range(r.Start, r.End - 1).InsertAfter(f'，{size}')

    def outline_demote(self, demote_level):
        """ 标题降级，降低level层 """
        for p in self.Paragraphs:
            p.Range.demote(demote_level)

    def set_style(self, obj, name):
        """ 给Paragraph、Range等按名称设置样式

        :param obj: 当前doc下某个含有Style成员属性的子对象
        :param name: 样式名称
        """
        setattr(obj, 'Style', self.Styles(name))

    def add_paragraph(self, text='', style=None):
        """ 自定义的插入段落

        默认的插入比较麻烦，新建段落、插入文本、设置格式要多步实现，这里封装支持在一步进行多种操作
        """
        p = self.Paragraphs.Add()
        if text:
            p.Range.InsertBefore(text)
        if style:
            p.Style = self.Styles(style)
        return p


class XlWin32WordRange:
    """ range是以下标0开始，左闭右开的区间

    当一个区间出现混合属性，比如有的有加粗，有的没加粗时，标记值为 app.wd('Undefined') 9999999
    vba的True是值-1，False是值0
    """

    def set_hyperlink(self, url):
        """ 给当前rng添加超链接
        """
        doc = self.Parent
        doc.Hyperlinks.Add(self, url)

    @property
    def chars(self):
        # 有特殊换行，ch.Text可能会得到 '\r\x07'，为了位置对应，只记录一个字符
        return ''.join([ch.Text[0] for ch in self.Characters])

    def char_range(self, start=0, end=None):
        """ 定位rng中的子range对象，这里是以可见字符Characters计数的

        :param start: 下标类似切片的规则
        :param end: 见start描述，允许越界，允许负数
            默认不输入表示匹配到末尾
        """
        n = self.Characters.Count
        if end is None or end > n:
            end = n
        elif end < 0:
            end = n + end
        start_idx, end_idx = self.Characters(start + 1).Start, self.Characters(end).End
        return self.Document.Range(start_idx, end_idx)

    def shifting(self, left=0, right=0):
        """ range左右两边增加偏移量，返回重定位的rng

        常用语段落定位，要在段落末尾增加内容时
        >> rng2 = p.Range.shifting(right=-1)
        """
        return self.Document.Range(self.Start + left, self.End + right)

    def demote(self, demote_level):
        """ 标题降级，降低level层 """
        name = self.Style.NameLocal  # 获得样式名称
        m = re.match(r'标题 (\d)$', name)
        if m:
            lvl = int(m.group(1))
            new_lvl = lvl + demote_level
            new_style = f'标题 {new_lvl}' if new_lvl < 10 else '正文'
            self.Style = self.Parent.Styles(new_style)

    def set_font(self, font_fmt):
        """ 设置字体各种格式

        :param dict|str font_fmt:
            dict，定制格式
                布尔类型：Bold、Italic、Subscript、Superscript
                可布尔的值类型：Underline
                    支持的格式见：https://docs.microsoft.com/en-us/office/vba/api/word.wdunderline
                值类型：Name、Size、Color、UnderlineColor
            str，使用现有样式名
        """
        if isinstance(font_fmt, dict):
            font = self.Font
            for k, v in font_fmt.items():
                setattr(font, k, v)
        elif isinstance(font_fmt, str):
            self.Style = self.Parent.Styles(font_fmt)
        else:
            raise ValueError

    def insert_before(self, text, font_fmt=None):
        """ 对原InsertBefore的功能封装

        :return: 增加返回值，是新插入内容的range定位
        """
        start1, end1 = self.Start, self.End
        self.InsertBefore(text)
        bias = self.End - end1  # 新插入的内容大小
        new_rng = self.Document.Range(start1, start1 + bias)
        if font_fmt:
            new_rng.set_font(font_fmt)
        return new_rng

    def insert_after(self, text, font_fmt=None):
        """ 同insert_before，是InsertAfter的重封装
        """
        # 1
        start1, end1 = self.Start, self.End

        # 2 往后插入，会排除\r情况
        doc = self.Document
        ch = doc.Range(end1 - 1, end1).Text
        if ch == '\r':
            end1 -= 1
            self = doc.Range(start1, end1)

        # 3
        self.InsertAfter(text)
        bias = self.End - end1
        new_rng = self.Document.Range(end1, end1 + bias)
        if font_fmt:
            new_rng.set_font(font_fmt)
        return new_rng


class XlWin32WordHyperlink:
    @property
    def netloc(self):
        from urllib.parse import urlparse
        linkp = urlparse(self.Name)  # 链接格式解析
        # netloc = linkp.netloc or Path(linkp.path).name
        netloc = linkp.netloc or linkp.scheme  # 可能是本地文件，此时记录其所在磁盘
        return netloc

    @property
    def name(self):
        """ 这个是转成明文的完整链接，如果要编码过的，可以取link.Name """
        from urllib.parse import unquote
        return unquote(self.Name)


def rebuild_document_by_word(fmt='html', translate=False, navigation=False, visible=False, quit=None):
    """ 将剪切板的内容粘贴到word重新排版，再转成fmt格式的文档，用浏览器打开

    这个功能只能在windows平台使用，并且必须要安装有Word软件。

    一般用于英文网站，生成双语阅读的模板，再调用谷歌翻译。
    生成的文档如果有需要，一般是手动复制整理到另一个docx文件中。

    Args:
        fmt: 输出文件类型
            常见的可以简写：html、pdf、txt
            其他特殊需求可以用word原变量名：wdFormatDocument
        visible: 是否展示运行过程，如果不展示，默认最后会close文档
        quit: 运行完是否退出应用
        translate: html专用业务功能，表示是否对p拷贝一份notranslate的对象，用于谷歌翻译双语对照
        navigation: 是否增加导航栏
            注意，使用导航栏后，页面就无法直接使用谷歌翻译了
            但可以自己进入_content文件，使用谷歌翻译处理，自覆盖保存
            然后再回到_index文件，刷新即可
    """
    import pyperclip
    from pyxllib.text.xmllib import BeautifulSoup, html_bitran_template, MakeHtmlNavigation

    # 1 保存的临时文件名采用etag
    f = File(get_etag(pyperclip.paste()), Dir.TEMP, suffix=fmt)
    app = XlWin32WordApplication.get_app(visible=visible)
    app.check_close(f)
    doc = app.new_doc(f)
    doc.Activate()
    app.Selection.Paste()

    # 2 如果需要，也可以在这个阶段，插入word自动化的操作，而不是后续在html层面操作
    # 统计每节内容长度，每个字母1B，每个汉字2B
    doc.add_section_size()
    file = doc.save(f, fmt)
    doc.Close()

    # 3 html格式扩展功能
    if fmt == 'html':
        # 3.1 默认扩展功能
        s = file.read(encoding='gbk')
        # s = s.replace('\xa0', '')  # 不知道这样去除\xa0行不行，等下次遇到尝试
        bs = BeautifulSoup(s, 'lxml')
        bs.head_add_number()  # 给标题加上编号
        # bs.head_add_size()  # 显示每节内容长短
        content = str(bs)

        # TODO 识别微信、pydoc，然后做一些自动化清理？
        # TODO 过度缩进问题？

        # 3.2 双语对照阅读
        if translate:
            # word生成的html固定是gbk编码
            content = html_bitran_template(content)

        # 原文是gbk，但到谷歌默认是utf8，所以改一改
        # 这样改后问题是word可能又反而有问题了，不过word本来只是跳板，并不是要用word编辑html
        file.write(content, encoding='utf8')

        # 3.3 导航栏功能
        # 作为临时使用可以开，如果要复制到word，并没有必要
        if navigation:
            file = MakeHtmlNavigation.from_file(file, encoding='utf8', number=False, text_catalogue=False)

    if quit:
        app.Quit()

    return file
